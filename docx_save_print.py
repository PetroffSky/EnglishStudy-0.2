import docx
from docx import Document
from docx.shared import Pt, Cm, Mm
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
import datetime as dt
import subprocess


class Save_Print:
    def __init__(self):
        self.save_data_to_docx()


    def save_data_to_docx(self):
        time_stamp = f"{dt.datetime.now().date()}_{dt.datetime.now().hour}-{dt.datetime.now().minute}"
        self.filename = f"Results/EnglishStudy_{time_stamp}.docx"
        self.doc = docx.Document()
        # доступ к первой секции:
        section = self.doc.sections[0]
        # высота листа в сантиметрах
        section.page_height = Mm(210)
        # ширина листа в сантиметрах
        section.page_width = Mm(148)
        # левое поле в миллиметрах
        section.left_margin = Mm(10)
        # правое поле в миллиметрах
        section.right_margin = Mm(10)
        # верхнее поле в миллиметрах
        section.top_margin = Mm(10)
        # нижнее поле в миллиметрах
        section.bottom_margin = Mm(5)

        with open('Results/translate.txt', 'r', encoding='utf-8') as translate:
            for stroke in translate.readlines():
                if len(stroke) > 0:
                    word, translate, transcription, url = stroke.split(',')
                    print(word, translate, transcription)

                    self.doc_formatter(f"Слово: {word}", 32)
                    self.image_to_doc(f"Results/Images/{word}.jpg", 6)
                    self.doc_formatter("Пишется так: ", 22)
                    self.doc_formatter(f"{translate}", 32)
                    self.doc_formatter("Произносится как: ", 22)
                    self.doc_formatter(f"{transcription}", 32)
                    self.image_to_doc(f"Results/Images/QR_{word}.png", 2)

                    p = self.doc.add_paragraph()
                    p.add_run().add_break(WD_BREAK.PAGE)

            self.doc.save(self.filename)
        print(f"Формирование документа {self.filename} готово!")
        self.print_docx()

    def doc_formatter(self, text, font_size):
        run = self.doc.add_paragraph()
        run.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = run.add_run(text)
        font = run.font
        font.name = 'Calibri'
        font.size = Pt(font_size)

    def image_to_doc(self, path, width):
        run = self.doc.add_paragraph()
        run.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = run.add_run("")
        run.add_picture(path, width=Cm(width), height=Cm(width))


    def print_docx(self):
        print("Открываю конечный файл для просмотра!")
        program_path = r'C:\Program Files\Microsoft Office 15\root\office15\WINWORD.EXE'
        subprocess.Popen([program_path, self.filename], shell=True)
        # Новая функция: запрос - отправить файл на печать
        # os.startfile(r'result.docx', 'print')


if __name__ == "__main__":
    Save_Print()